from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "BRANDING_AND_COMMUNITY_REPORT.md"
OUTPUT = ROOT / "Rocci_Branding_and_Community_Foundation.docx"

INK = "201E1C"
MUTED = "6D6862"
FAINT = "918A82"
CANVAS = "FAF9F6"
SURFACE = "FFFFFF"
CORAL = "B92F19"
FOLD_CORAL = "E64B2F"
SOFT_CORAL = "FFF0E9"
PLUM = "6D4AFF"
LINE = "D7D0C7"
CODE_BG = "282522"
CODE_INK = "F8F3ED"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_numbering_start(doc: Document, paragraph, start: int) -> None:
    """Attach a real decimal numbering definition with an explicit start."""
    numbering = doc.part.numbering_part.element
    style_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    abstract_num_id = None
    for num in numbering.findall(qn("w:num")):
        if int(num.get(qn("w:numId"))) == style_num_id:
            abstract_num_id = int(num.find(qn("w:abstractNumId")).get(qn("w:val")))
            break
    if abstract_num_id is None:
        raise RuntimeError("List Number style has no numbering definition")

    num_ids = [int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    num.append(override)
    numbering.append(num)

    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, color=CORAL) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, color_node, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\))")


def add_inline(paragraph, text: str, *, base_color=INK, base_size=10.5) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=INK, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Menlo", size=9.2, color=CORAL)
            r_pr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), SOFT_CORAL)
            r_pr.append(shd)
        else:
            link = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            add_hyperlink(paragraph, link.group(1), link.group(2))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=base_color)


def paragraph_border(paragraph, *, left=None, bottom=None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    if left:
        node = OxmlElement("w:left")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(left[0]))
        node.set(qn("w:space"), str(left[1]))
        node.set(qn("w:color"), left[2])
        borders.append(node)
    if bottom:
        node = OxmlElement("w:bottom")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(bottom[0]))
        node.set(qn("w:space"), str(bottom[1]))
        node.set(qn("w:color"), bottom[2])
        borders.append(node)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    set_run_font(run, size=8.5, color=FAINT)


def set_latest_image_alt(doc: Document, description: str) -> None:
    doc_pr = doc.inline_shapes[-1]._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, CORAL),
        ("Heading 2", 13, 14, 10, CORAL),
        ("Heading 3", 12, 10, 5, INK),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Code Block"]
    style.font.name = "Menlo"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Menlo")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Menlo")
    style.font.size = Pt(8.8)
    style.font.color.rgb = RGBColor.from_string(CODE_INK)
    style.paragraph_format.left_indent = Inches(0.16)
    style.paragraph_format.right_indent = Inches(0.16)
    style.paragraph_format.space_before = Pt(5)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    if "Figure Caption" not in styles:
        style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Figure Caption"]
    style.font.name = "Calibri"
    style.font.size = Pt(8.5)
    style.font.italic = True
    style.font.color.rgb = RGBColor.from_string(MUTED)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.keep_with_next = False


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.text = "ROCCI  /  EXPLORATORY REPORT"
    set_run_font(header.runs[0], size=8, color=FAINT, bold=True)
    header.paragraph_format.space_after = Pt(0)
    paragraph_border(header, bottom=(6, 5, LINE))

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(44)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("BRAND ARCHITECTURE · VISUAL IDENTITY · COMMUNITY")
    set_run_font(run, size=8.5, color=CORAL, bold=True)

    reset = doc.add_paragraph()
    reset.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reset.paragraph_format.space_after = Pt(20)
    run = reset.add_run("ZERO-BASED IDENTITY EXPLORATION")
    set_run_font(run, size=10, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Rocci branding and\ncommunity foundation")
    set_run_font(run, size=27, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("A reversible public-preview recommendation for an open-source Roc ecosystem project")
    set_run_font(run, size=12.5, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    paragraph_border(rule, bottom=(16, 2, FOLD_CORAL))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("17 August 2026  ·  Exploratory  ·  Community review required")
    set_run_font(run, size=9.5, color=MUTED, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Not trademark clearance. Not authorization to use or modify the Roc logo.")
    set_run_font(run, size=8.5, color=FAINT, italic=True)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.top_margin = Inches(1)
    body.right_margin = Inches(1)
    body.bottom_margin = Inches(1)
    body.left_margin = Inches(1)
    body.header_distance = Inches(0.492)
    body.footer_distance = Inches(0.492)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    add_inline(p, text, base_color=INK, base_size=11)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), SOFT_CORAL)
    p_pr.append(shd)
    paragraph_border(p, left=(18, 7, FOLD_CORAL))


def table_widths(headers: list[str]) -> list[int]:
    n = len(headers)
    if n == 2:
        return [2700, 6660]
    if n == 3:
        return [2000, 2600, 4760]
    if n == 4 and headers[0] == "Role":
        return [1700, 1500, 1500, 4660]
    if n == 4:
        return [1400, 2200, 3300, 2460]
    return [9360 // n] * n


def add_table(doc: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = table_widths(headers)
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, SOFT_CORAL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, value, base_color=INK, base_size=9)
        for run in p.runs:
            run.bold = True
    set_repeat_table_header(table.rows[0])
    keep_table_row_together(table.rows[0])
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value, base_color=INK, base_size=8.7)
        keep_table_row_together(table.rows[-1])
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def normalize_table_row(line: str) -> list[str]:
    return [value.strip() for value in line.strip().strip("|").split("|")]


def parse_report(doc: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    skip_title = True
    page_break_headings = {
        "Logo and icon exploration",
        "Community foundation and public-preview plan",
        "Sources",
    }
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, name="Menlo", size=8.8, color=CODE_INK)
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), CODE_BG)
                p_pr.append(shd)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if skip_title and line.startswith("# "):
            skip_title = False
            i += 1
            continue
        if line.startswith("**Exploratory recommendation"):
            i += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if heading in page_break_headings and doc.paragraphs[-1].text:
                doc.add_page_break()
            p = doc.add_paragraph(heading, style="Heading 1")
            if heading == "Executive recommendation":
                p.paragraph_format.space_before = Pt(0)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
            i += 1
            continue
        if line.startswith("| ") and i + 1 < len(lines) and re.match(r"^\|[ :\-\|]+\|$", lines[i + 1]):
            rows = [normalize_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(normalize_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue
        image_match = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", line)
        if image_match:
            image_path = ROOT / image_match.group(2)
            doc.add_picture(str(image_path), width=Inches(4.65))
            set_latest_image_alt(doc, image_match.group(1) + " generated Rocci logo concept sheet.")
            p = doc.paragraphs[-1]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            cap = doc.add_paragraph(style="Figure Caption")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(image_match.group(1) + " — generated concept probe; not a production logo")
            i += 1
            continue
        if line.startswith("> "):
            quote_parts = [line[2:].strip()]
            i += 1
            while i < len(lines) and lines[i].startswith("> "):
                quote_parts.append(lines[i][2:].strip())
                i += 1
            add_callout(doc, " ".join(quote_parts))
            continue
        bullet_match = re.match(r"^(\s*)- (.+)$", line)
        if bullet_match:
            base_indent = len(bullet_match.group(1))
            parts = [bullet_match.group(2).strip()]
            i += 1
            while i < len(lines) and lines[i].strip():
                next_indent = len(lines[i]) - len(lines[i].lstrip())
                next_text = lines[i].strip()
                if next_indent <= base_indent or re.match(r"^(?:- |\d+\. )", next_text):
                    break
                parts.append(next_text)
                i += 1
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375 + min(base_indent, 6) * 0.055)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, "\u00a0" + " ".join(parts))
            continue
        numbered = re.match(r"^(\s*)(\d+)\. (.+)$", line)
        if numbered:
            base_indent = len(numbered.group(1))
            parts = [numbered.group(3).strip()]
            i += 1
            while i < len(lines) and lines[i].strip():
                next_indent = len(lines[i]) - len(lines[i].lstrip())
                next_text = lines[i].strip()
                if next_indent <= base_indent or re.match(r"^(?:- |\d+\. )", next_text):
                    break
                parts.append(next_text)
                i += 1
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.375 + min(base_indent, 6) * 0.055)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            set_numbering_start(doc, p, int(numbered.group(2)))
            add_inline(p, " ".join(parts))
            continue
        if not line.strip():
            i += 1
            continue

        parts = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip() or nxt.startswith(("#", "- ", "> ", "```", "|", "![")) or re.match(r"^\d+\. ", nxt):
                break
            parts.append(nxt.strip())
            i += 1
        text = " ".join(parts)
        if text.startswith("This report prepares"):
            add_callout(doc, text)
        else:
            p = doc.add_paragraph()
            add_inline(p, text)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)
    add_cover(doc)
    parse_report(doc)

    props = doc.core_properties
    props.title = "Rocci branding and community foundation"
    props.subject = "Exploratory brand architecture, visual identity, SEO, and public-preview community plan"
    props.author = "Rocci Project"
    props.keywords = "Rocci, Rocdown, Rocs, Roc, Datastar, branding, community, open source"
    props.comments = "Generated from the repository branding report; visual concepts are exploratory."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
